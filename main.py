import barcode
import csv
from barcode.writer import ImageWriter
from PIL import Image
import os, re
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
import datetime
from docx.shared import Pt


def generateEAN(file):

    print("    "+file)
    generatePNG(file)
    document = Document("VZOR-EAN.docx")

    table = document.tables[0]

    for sRow in table.rows:
        for sCell in sRow.cells:
            sCell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraph = sCell.paragraphs[0]
            run = paragraph.add_run()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run.add_picture("cache/"+file+".png", width=Cm(4.74), height=Cm(2.41))
    #
    document.save(folder+"/"+file+".docx")


def generateNames(file, text, count):
    print("  "+file)
    print("    "+text)

    document = Document("VZOR-POPISEK.docx")

    table = document.tables[0]
    full_text = file + "\n " + text + "Ks/kt " + count + " x " + file
    for sRow in table.rows:
        for sCell in sRow.cells:
            sCell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraph = sCell.paragraphs[0]
            run = paragraph.add_run()
            run.bold = True
            run.font.size = Pt(11)
            run.text = full_text
    #
    document.save(folder+"/POPISEK "+file+".docx")


def generatePNG(number):
    # Generate barcode
    ean = barcode.get('ean13', number, writer=ImageWriter())

    filename = ean.save("cache/"+number, {"module_width": 0.4, "module_height": 13, "font_size": 20, "text_distance": 1,
                                          "quiet_zone": 2})
    # Resize
    image = Image.open(filename)
    resized = image.resize((226, 100))
    resized.save(filename)


def clearCache():
    print("\nMažu cache..")
    for filename in os.listdir("cache"):
        os.remove("cache/"+filename)
    os.rmdir("cache")


# Main
if __name__ == "__main__":
    fields = []
    rows = []
    with open('togen.csv', 'r', encoding="UTF8") as csvfile:
        csvreader = csv.reader(csvfile, delimiter=';')

        fields = next(csvreader)
        for row in csvreader:
            # remove non-numeric characters
            row[0] = re.sub(r'[^0-9]', '', row[0])
            row[1] = re.sub(r'[^0-9]', '', row[1])
            row[3] = re.sub(r'[^0-9]', '', row[3])
            if row[0] != "" and row[1] != "" and row[2] != "" and row[3] != "":
                rows.append(row)

    folder = datetime.datetime.now().strftime("%02d.%02m.%04Y")
    if not (os.path.exists(folder)):
        os.mkdir(folder)
    else:
        for filename in os.listdir(folder):
            os.remove(folder+"/"+filename)
    if not (os.path.exists("cache")):
        os.mkdir("cache")

    print("Generuji data..")
    for row in rows:
        generateNames(row[0], row[2], row[3])
        generateEAN(row[1])

    clearCache()
    print("  Hotovo!")
